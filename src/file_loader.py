import os
from pathlib import Path

from langchain_community.document_loaders.csv_loader import CSVLoader
from langchain.document_loaders.base import BaseLoader
from langchain.schema import Document
from langchain_community.document_loaders.unstructured import UnstructuredFileLoader
from docx import Document as DocxDocument

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredHTMLLoader,
    UnstructuredMarkdownLoader,
    UnstructuredPowerPointLoader,
    UnstructuredWordDocumentLoader,
    WebBaseLoader,
    Docx2txtLoader,
    UnstructuredExcelLoader,
)

# Create file and loader mapping according to different extensions
FILE_LOADER_MAPPING = {
    ".csv": (CSVLoader, {"encoding": "utf-8"}),
    ".xlsx": (UnstructuredExcelLoader, {}),
    ".doc": (UnstructuredWordDocumentLoader, {}),
    ".docx": (Docx2txtLoader, {}),
    ".html": (UnstructuredHTMLLoader, {}),
    ".md": (UnstructuredMarkdownLoader, {}),
    ".pdf": (PyPDFLoader, {}),
    ".ppt": (UnstructuredPowerPointLoader, {}),
    ".pptx": (UnstructuredPowerPointLoader, {}),
    ".txt": (TextLoader, {}),
    # You can add more mappings for other file extensions and loaders as needed
}


def fallback_docx_loader(file_path):
    """
    A fallback loader for .docx files using python-docx.
    """
    try:
        doc = DocxDocument(file_path)
        content = "\n".join([p.text for p in doc.paragraphs])
        return [Document(page_content=content, metadata={"source": file_path})]
    except Exception as e:
        print(f"Fallback loader failed for {file_path}: {e}")
        return []


def load_document(
    folder_path: str,
    mapping: dict = FILE_LOADER_MAPPING,
    default_loader: BaseLoader = UnstructuredFileLoader,
) -> list[Document]:
    loaded_documents = []

    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)  # Full file path
        ext = "." + file_path.rsplit(".", 1)[-1].lower()  # File extension

        if ext in mapping:
            loader_class, loader_args = mapping[ext]
            try:
                loader = loader_class(file_path, **loader_args)
                docs = loader.load()
                # Add file path as metadata to each document
                for doc in docs:
                    doc.metadata["source"] = filename
                loaded_documents.extend(docs)
            except Exception as e:
                print(f"Error loading file {file_path} with {loader_class}: {e}")

                # Use fallback loader for .docx files if primary loader fails
                if ext == ".docx":
                    loaded_documents.extend(fallback_docx_loader(file_path))

        else:
            try:
                loader = default_loader(file_path)
                docs = loader.load()
                # Add file path as metadata to each document
                for doc in docs:
                    doc.metadata["source"] = file_path  # Add source metadata
                    print(f"Document loaded with metadata: {doc.metadata}")  # Debug log

                loaded_documents.extend(docs)
            except Exception as e:
                print(f"Error loading file {file_path} with default loader: {e}")

    return loaded_documents


# Example usage
# documents = load_document(folder_path="test")
# print(documents)